"""
Serviço para geração de contratos a partir de template Word
"""
import os
import tempfile
import datetime
from docx import Document
from typing import Dict, Any, Optional, List
import logging
import json
from fuzzywuzzy import fuzz, process

logger = logging.getLogger(__name__)

class ContractGenerationService:
    """Serviço responsável pela geração de contratos usando template Word"""
    
    # Mapeamento das empresas BPO
    BPO_COMPANIES_MAPPING = {
        'HR HILL': {
            'razao_social': 'Hr Hill Servicos Administrativos Ltda.',
            'cnpj': '36.446.561/0001-66'
        },
        'JMW SERVIÇOS': {
            'razao_social': 'JMW Servicos Administrativos Ltda.',
            'cnpj': '36.448.288/0001-09'
        },
        'GF SERVIÇOS': {
            'razao_social': 'GF Servicos De Contabilidade Ltda.',
            'cnpj': '36.583.021/0001-24'
        },
        'ACARNEGIE SERVIÇOS': {
            'razao_social': 'Acarnegie Servicos Administrativos Ltda.',
            'cnpj': '40.601.894/0001-90'
        },
        'E.REEVE SERVIÇOS': {
            'razao_social': 'E. Reeve Musk Servicos de Contabilidade Ltda.',
            'cnpj': '40.897.585/0001-09'
        },
        'S. JOBS SERVIÇOS': {
            'razao_social': 'S.Jobs Servicos de Contabilidade Ltda.',
            'cnpj': '40.933.869/0001-03'
        },
        'GF FINANCE': {
            'razao_social': 'GF Finance Ltda.',
            'cnpj': '44.613.528/0001-01'
        },
        'GF PAYROLL': {
            'razao_social': 'GF Payroll Ltda.',
            'cnpj': '44.612.639/0001-01'
        },
        'GF ACCOUNTING': {
            'razao_social': 'GF Accounting Ltda.',
            'cnpj': '44.615.004/0001-50'
        },
        'GF TAX': {
            'razao_social': 'GF Tax Servicos de Contabilidade Ltda.',
            'cnpj': '44.573.718/0001-42'
        },
        'GF LEGAL': {
            'razao_social': 'GF Legal Ltda.',
            'cnpj': '40.591.977/0001-45'
        }
    }
    
    @staticmethod
    def get_template_path(contract_type: str = 'bpo_contabil_completo') -> str:
        """
        Retorna o caminho do template baseado no tipo de contrato
        
        Args:
            contract_type: Tipo do contrato (ex: 'bpo_contabil_completo', 'bpo_contabil_completo_bicolunado')
            
        Returns:
            Caminho para o template .docx correspondente
        """
        template_filename = f"{contract_type}.docx"
        template_path = os.path.join(os.path.dirname(__file__), '..', 'utils', template_filename)
        
        if not os.path.exists(template_path):
            # Fallback para o template padrão
            logger.warning(f"Template {template_filename} não encontrado, usando template padrão")
            template_path = os.path.join(os.path.dirname(__file__), '..', 'utils', 'bpo_contabil_completo.docx')
        
        return os.path.abspath(template_path)
    
    def __init__(self, template_path: str = None, contract_type: str = 'bpo_contabil_completo'):
        """
        Inicializa o serviço com o caminho do template
        
        Args:
            template_path: Caminho para o arquivo template .docx (opcional)
            contract_type: Tipo do contrato para auto-resolver template (opcional)
        """
        if template_path is None:
            # Resolver template baseado no tipo
            template_path = self.get_template_path(contract_type)
        
        self.template_path = os.path.abspath(template_path)
        
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Template não encontrado: {self.template_path}")
    
    def _find_bpo_company_by_fuzzy_match(self, bpo_name: str) -> Optional[Dict[str, str]]:
        """
        Encontra a empresa BPO usando fuzzy matching
        
        Args:
            bpo_name: Nome de referência da empresa BPO
            
        Returns:
            Dicionário com razao_social e cnpj da empresa, ou None se não encontrado
        """
        if not bpo_name or bpo_name.strip() == '':
            return None
            
        # Limpar o nome de entrada
        bpo_name_clean = bpo_name.strip().upper()
        
        # Lista de chaves para busca
        company_keys = list(self.BPO_COMPANIES_MAPPING.keys())
        
        # Fazer fuzzy matching
        best_match, score = process.extractOne(bpo_name_clean, company_keys)
        
        # Usar threshold de 70% de similaridade
        if score >= 70:
            logger.info(f"Fuzzy match encontrado: '{bpo_name}' -> '{best_match}' (score: {score})")
            return self.BPO_COMPANIES_MAPPING[best_match]
        else:
            logger.warning(f"Nenhum match encontrado para '{bpo_name}' (melhor: '{best_match}' com score {score})")
            return None
    
    def _generate_signature_section(self, bpo_companies: List[Dict[str, str]]) -> None:
        """
        Gera a seção de assinaturas dinamicamente baseada nas empresas BPO
        
        Args:
            bpo_companies: Lista de empresas BPO com nome e CNPJ
        """
        try:
            # Carregar template
            doc = Document(self.template_path)
            
            # Mapear parágrafos de assinatura por nome da empresa
            # Baseado na estrutura encontrada: S.JOBS (17), E. REEVE MUSK (20), GF PAYROLL (25)
            signature_paragraphs = {}
            
            # Encontrar os parágrafos das empresas de assinatura
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text == 'S.JOBS.':
                    signature_paragraphs['S.JOBS'] = i
                elif text == 'E. REEVE MUSK SERVIÇOS DE CONTABILIDADE LTDA.':
                    signature_paragraphs['E.REEVE'] = i
                elif text == 'GF PAYROLL.':
                    signature_paragraphs['GF PAYROLL'] = i
                    
            logger.info(f"Parágrafos de assinatura encontrados: {signature_paragraphs}")
            
            # Limpar todos os parágrafos de assinatura existentes
            for company_key, paragraph_index in signature_paragraphs.items():
                doc.paragraphs[paragraph_index].text = ""
                logger.info(f"Limpado parágrafo {paragraph_index} da empresa {company_key}")
            
            # Adicionar novas empresas nos parágrafos corretos
            available_positions = sorted(signature_paragraphs.values())
            
            for i, bpo_company in enumerate(bpo_companies):
                if i < len(available_positions):
                    paragraph_index = available_positions[i]
                    company_name = bpo_company['nome'].upper()
                    
                    # Simplificar o nome para assinatura (remover LTDA e deixar mais curto)
                    signature_name = company_name
                    if 'E. REEVE MUSK' in company_name:
                        signature_name = 'E. REEVE MUSK SERVIÇOS DE CONTABILIDADE LTDA.'
                    elif 'GF SERVICOS DE CONTABILIDADE' in company_name:
                        signature_name = 'GF SERVIÇOS.'
                    elif 'HR HILL' in company_name:
                        signature_name = 'HR HILL.'
                    elif 'S.JOBS' in company_name:
                        signature_name = 'S.JOBS.'
                    elif 'GF PAYROLL' in company_name:
                        signature_name = 'GF PAYROLL.'
                    else:
                        # Para outras empresas, usar nome simplificado
                        signature_name = company_name.replace(' LTDA.', '.').replace(' LTDA', '.')
                        if not signature_name.endswith('.'):
                            signature_name += '.'
                    
                    doc.paragraphs[paragraph_index].text = signature_name
                    logger.info(f"Adicionado '{signature_name}' no parágrafo {paragraph_index}")
            
            # Salvar template modificado temporariamente
            temp_template_path = self.template_path.replace('.docx', '_temp.docx')
            doc.save(temp_template_path)
            self.template_path = temp_template_path
            
            logger.info(f"Seção de assinaturas atualizada com {len(bpo_companies)} empresas")
            
        except Exception as e:
            logger.error(f"Erro ao gerar seção de assinaturas: {str(e)}")
            # Não falhar por conta das assinaturas, continuar sem elas
            logger.warning("Continuando geração sem modificar assinaturas")

    def _extract_bpo_companies(self, company_data: Dict[str, Any]) -> tuple[str, str, List[Dict[str, str]]]:
        """
        Extrai dinamicamente as empresas BPO dos dados da empresa e gera o texto formatado
        
        Args:
            company_data: Dados da empresa contendo campos BPO
            
        Returns:
            Tupla com: (texto_portugues, texto_ingles, lista_empresas)
        """
        logger.info(f"=== Iniciando extração de empresas BPO ===")
        logger.info(f"Dados recebidos - campos disponíveis: {list(company_data.keys())}")
        logger.info(f"Amostra dos dados: {dict(list(company_data.items())[:5])}")  # Primeiros 5 campos
        
        try:
            # Verificar se temos dados completos da empresa ou apenas dados básicos
            basic_fields = {'razao_social', 'cnpj', 'endereco'}
            has_only_basic_data = (
                len(set(company_data.keys()) - basic_fields) == 0 or
                all(key in basic_fields for key in company_data.keys() if company_data.get(key))
            )
            
            if has_only_basic_data:
                logger.warning("Detectados apenas dados básicos da empresa, usando empresas BPO padrão")
                bpo_companies = {"GF SERVIÇOS", "E.REEVE SERVIÇOS", "HR HILL"}
            else:
                # Campos BPO para verificar
                bpo_fields = [
                    'BPO Contábil Faturado',
                    'BPO Fiscal Faturado', 
                    'BPO Folha Faturado',
                    'BPO Financeiro Faturado',
                    'BPO RH Faturado',
                    'BPO Legal Faturado',
                    'Diversos In. Faturado',
                    'Implantação Faturado'
                ]
                
            # Extrair valores únicos e não nulos dos campos BPO
            bpo_companies = set()
            
            # Se company_data tem um campo 'companie_data' (JSON), usar ele
            data_source = company_data.get('companie_data', {})
            if isinstance(data_source, str):
                try:
                    data_source = json.loads(data_source)
                except json.JSONDecodeError:
                    data_source = {}
            
            # Se não encontrou no companie_data, usar os dados diretos
            if not data_source:
                data_source = company_data
            
            logger.info(f"Buscando empresas BPO nos dados: {list(data_source.keys())[:10]}...")  # Log primeiro 10 campos
            
            for field in bpo_fields:
                value = data_source.get(field)
                if value and value.strip() and value.lower() not in ['null', 'none', '']:
                    # Normalizar o nome para evitar duplicatas
                    normalized_name = value.strip().replace('.', '').replace(' ', '').upper()
                    # Usar o valor original como chave, mas evitar duplicatas baseado no nome normalizado
                    found_duplicate = False
                    for existing_company in bpo_companies:
                        existing_normalized = existing_company.replace('.', '').replace(' ', '').upper()
                        if existing_normalized == normalized_name:
                            found_duplicate = True
                            break
                    
                    if not found_duplicate:
                        bpo_companies.add(value.strip())
                        logger.info(f"Encontrada empresa BPO: {field} = {value}")
                    else:
                        logger.info(f"Empresa BPO duplicada ignorada: {field} = {value}")
            
            # Se não encontrou nenhuma empresa BPO, usar um padrão
            if not bpo_companies:
                logger.warning("Nenhuma empresa BPO encontrada nos dados, usando empresas padrão")
                bpo_companies = {"GF SERVIÇOS", "E.REEVE SERVIÇOS", "HR HILL"}
            
            # Endereço padrão para todas as empresas
            endereco_padrao = "Av. Dr. Cardoso de Melo, nº 1608, 8º andar, 81-B, Vila Olímpia, São Paulo/SP, CEP: 04548-005"
            endereco_padrao_en = "Av. Dr. Cardoso de Melo, No. 1608, 8th floor, suite 81-B, Vila Olímpia, São Paulo/SP, ZIP: 04548-005"
            
            # Gerar texto formatado usando fuzzy matching
            companies_list = sorted(list(bpo_companies))
            logger.info(f"Lista de empresas BPO ordenada: {companies_list}")
            letters = ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j']  # Suporte até 10 empresas
            
            formatted_companies_pt = []
            formatted_companies_en = []
            companies_for_signature = []  # Lista de dicionários para assinaturas
            
            for i, bpo_name in enumerate(companies_list):
                if i < len(letters):
                    letter = letters[i]
                    
                    # Buscar dados da empresa usando fuzzy matching
                    empresa_data = self._find_bpo_company_by_fuzzy_match(bpo_name)
                    logger.info(f"Processando empresa {i+1}/{len(companies_list)}: '{bpo_name}' -> {empresa_data}")
                    
                    if empresa_data:
                        razao_social = empresa_data['razao_social']
                        cnpj = empresa_data['cnpj']
                    else:
                        # Fallback caso não encontre match
                        razao_social = f"{bpo_name} LTDA."
                        cnpj = "00.000.000/0001-00"
                        logger.warning(f"Usando dados fallback para: {bpo_name}")
                    
                    # Adicionar à lista de assinaturas
                    companies_for_signature.append({
                        'nome': razao_social,
                        'cnpj': cnpj
                    })
                    
                    logger.info(f"Adicionada empresa para contrato: {razao_social} ({cnpj})")
                    
                    # Texto em português
                    formatted_text_pt = (
                        f"{letter}) {razao_social.upper()}, inscrita no CNPJ sob o nº {cnpj}, "
                        f"com sede à {endereco_padrao}"
                    )
                    formatted_companies_pt.append(formatted_text_pt)
                    
                    # Texto em inglês
                    formatted_text_en = (
                        f"{letter}) {razao_social.upper()}, enrolled with CNPJ under No. {cnpj}, "
                        f"with headquarters at {endereco_padrao_en}"
                    )
                    formatted_companies_en.append(formatted_text_en)
            
            # Função para juntar empresas com conjunção
            def join_companies(companies_list, conjunction):
                if len(companies_list) == 1:
                    return companies_list[0]
                elif len(companies_list) == 2:
                    return f"{companies_list[0]}; {conjunction} {companies_list[1]}"
                else:
                    return "; ".join(companies_list[:-1]) + f"; {conjunction} {companies_list[-1]}"
            
            # Juntar todas com "; " e adicionar "e"/"and" antes da última
            result_pt = join_companies(formatted_companies_pt, "e") + ", todas neste ato representadas na forma de seus atos constitutivos, doravante denominadas simplesmente como CONTRATADA;"
            result_en = join_companies(formatted_companies_en, "and") + ", all herein represented in accordance with their corporate documents, hereinafter referred to simply as the CONTRACTED PARTY;"
            
            logger.info(f"Texto BPO gerado com {len(companies_for_signature)} empresas")
            logger.info(f"Texto completo PT: {result_pt[:500]}...")  # Primeiros 500 caracteres
            logger.info(f"Texto completo EN: {result_en[:500]}...")  # Primeiros 500 caracteres
            
            return result_pt, result_en, companies_for_signature
            
        except Exception as e:
            logger.error(f"Erro ao extrair empresas BPO: {str(e)}")
            # Retornar um padrão em caso de erro usando dados reais
            endereco_padrao = "Av. Dr. Cardoso de Melo, nº 1608, 8º andar, 81-B, Vila Olímpia, São Paulo/SP, CEP: 04548-005"
            endereco_padrao_en = "Av. Dr. Cardoso de Melo, No. 1608, 8th floor, suite 81-B, Vila Olímpia, São Paulo/SP, ZIP: 04548-005"
            
            fallback_text_pt = (
                f"a) GF SERVICOS DE CONTABILIDADE LTDA., inscrita no CNPJ sob o nº 36.583.021/0001-24, com sede à {endereco_padrao}; "
                f"b) E. REEVE MUSK SERVIÇOS DE CONTABILIDADE LTDA., inscrita no CNPJ sob o nº 40.897.585/0001-09, com sede à {endereco_padrao}; "
                f"e c) HR HILL SERVIÇOS ADMINISTRATIVOS LTDA., inscrita no CNPJ sob o nº 36.446.561/0001-66, com sede à {endereco_padrao}, "
                f"todas neste ato representadas na forma de seus atos constitutivos, doravante denominadas simplesmente como CONTRATADA;"
            )
            
            fallback_text_en = (
                f"a) GF SERVICOS DE CONTABILIDADE LTDA., enrolled with CNPJ under No. 36.583.021/0001-24, with headquarters at {endereco_padrao_en}; "
                f"b) E. REEVE MUSK SERVIÇOS DE CONTABILIDADE LTDA., enrolled with CNPJ under No. 40.897.585/0001-09, with headquarters at {endereco_padrao_en}; "
                f"and c) HR HILL SERVIÇOS ADMINISTRATIVOS LTDA., enrolled with CNPJ under No. 36.446.561/0001-66, with headquarters at {endereco_padrao_en}, "
                f"all herein represented in accordance with their corporate documents, hereinafter referred to simply as the CONTRACTED PARTY;"
            )
            
            fallback_companies = [
                {"nome": "GF SERVIÇOS DE CONTABILIDADE LTDA.", "cnpj": "36.583.021/0001-24"},
                {"nome": "E. REEVE MUSK SERVIÇOS DE CONTABILIDADE LTDA.", "cnpj": "40.897.585/0001-09"},
                {"nome": "HR HILL SERVIÇOS ADMINISTRATIVOS LTDA.", "cnpj": "36.446.561/0001-66"}
            ]
            return fallback_text_pt, fallback_text_en, fallback_companies
    
    def generate_contract(self, company_data: Dict[str, Any]) -> str:
        """
        Gera um contrato preenchido com os dados da empresa
        
        Args:
            company_data: Dados da empresa contendo as informações necessárias
            
        Returns:
            Caminho para o arquivo temporário gerado
            
        Raises:
            ValueError: Se dados obrigatórios estiverem ausentes
            FileNotFoundError: Se template não existir
        """
        try:
            logger.info(f"Iniciando geração de contrato para empresa: {company_data.get('razao_social', 'N/A')}")
            logger.info(f"Usando template: {self.template_path}")
            
            # Validar dados obrigatórios
            required_fields = ['razao_social', 'cnpj', 'endereco']
            missing_fields = []
            
            for field in required_fields:
                if not company_data.get(field):
                    missing_fields.append(field)
            
            if missing_fields:
                raise ValueError(f"Campos obrigatórios ausentes: {', '.join(missing_fields)}")
            
            # Detectar se é um termo de distrato (qualquer variação)
            is_distrato = 'termo_distrato' in self.template_path.lower()
            is_distrato_sem_contrato = 'termo_distrato_sem_contrato' in self.template_path.lower()
            
            if is_distrato:
                # Para termos de distrato, também usar empresas BPO dinâmicas
                logger.info(f"Detectado termo de distrato - tipo: {'sem contrato' if is_distrato_sem_contrato else 'completo'}")
                
                # Gerar texto dinâmico das empresas BPO e lista para assinaturas
                empresas_bpo_texto_pt, empresas_bpo_texto_en, empresas_bpo_lista = self._extract_bpo_companies(company_data)
                
                # Carregar template
                doc = Document(self.template_path)
                
                # Primeiro, substituir o texto fixo das empresas BPO pelo dinâmico
                total_replacements = 0
                for i, paragraph in enumerate(doc.paragraphs):
                    original_text = paragraph.text
                    new_text = original_text
                    
                    # Detectar e substituir o texto fixo do GO FURTHER GROUP
                    if 'GO FURTHER GROUP, grupo de empresas independentes formado por:' in original_text and len(original_text) > 1000:
                        logger.info(f"Substituindo texto fixo das empresas BPO no parágrafo {i} do termo de distrato")
                        logger.info(f"Texto original tinha {len(original_text)} caracteres")
                        
                        # Substituir por texto dinâmico com formato de distrato
                        new_text = f"CONTRATADA: GO FURTHER GROUP, grupo de empresas independentes formado por: {empresas_bpo_texto_pt.replace('CONTRATADA: GO FURTHER GROUP, grupo de empresas independentes formado por: ', '').replace(', todas neste ato representadas na forma de seus atos constitutivos, doravante denominadas simplesmente como CONTRATADA;', ', todas neste ato representadas na forma de seus atos constitutivos.')}"
                        
                        total_replacements += 1
                        logger.info(f"Novo texto terá {len(new_text)} caracteres")
                        logger.info(f"Novo texto: {new_text[:200]}...")
                    
                    if new_text != original_text:
                        paragraph.text = new_text
                
                # Substituir placeholders específicos do cliente (apenas para distrato sem contrato)
                if is_distrato_sem_contrato:
                    field_mapping = {
                        'RAZÃO SOCIAL DA EMPRESA': company_data['razao_social'],
                        'CNPJ do Cliente': company_data['cnpj'],
                        'endereço completo': company_data['endereco'],
                        'RAZÃO SOCIAL DO CLIENTE': company_data['razao_social']
                    }
                    
                    logger.info(f"Campos de termo de distrato sem contrato a serem substituídos: {list(field_mapping.keys())}")
                    
                    # Substituir placeholders nos parágrafos
                    for i, paragraph in enumerate(doc.paragraphs):
                        original_text = paragraph.text
                        new_text = original_text
                        
                        for placeholder, value in field_mapping.items():
                            if placeholder in new_text:
                                new_text = new_text.replace(placeholder, str(value))
                                total_replacements += 1
                                logger.info(f"Substituído '{placeholder}' por '{value}' no parágrafo {i}")
                        
                        if new_text != original_text:
                            paragraph.text = new_text
                else:
                    # Para termo de distrato completo, usar placeholders similares aos contratos BPO
                    field_mapping = {
                        '[RAZÃO SOCIAL]': company_data['razao_social'],
                        '[CNPJ]': company_data['cnpj'],
                        '[ENDEREÇO]': company_data['endereco'],
                        '[COMPANY NAME]': company_data['razao_social'],
                        '[ADDRESS]': company_data['endereco']
                    }
                    
                    logger.info(f"Campos de termo de distrato completo a serem substituídos: {list(field_mapping.keys())}")
                    
                    # Substituir campos nos parágrafos
                    for i, paragraph in enumerate(doc.paragraphs):
                        original_text = paragraph.text
                        new_text = original_text
                        
                        for placeholder, value in field_mapping.items():
                            if placeholder in new_text:
                                new_text = new_text.replace(placeholder, str(value))
                                total_replacements += 1
                                logger.info(f"Substituído '{placeholder}' por '{value}' no parágrafo {i}")
                        
                        if new_text != original_text:
                            paragraph.text = new_text
                
                logger.info(f"Total de substituições realizadas no termo de distrato: {total_replacements}")
                
                # Gerar seção de assinaturas dinâmica para termo de distrato
                logger.info("Gerando seção de assinaturas dinâmica para termo de distrato")
                
                # Encontrar e substituir a seção de assinaturas
                signature_start = -1
                signature_end = -1
                
                for i, paragraph in enumerate(doc.paragraphs):
                    if 'PELA CONTRATADA:' in paragraph.text:
                        signature_start = i
                        logger.info(f"Início da seção de assinaturas encontrado no parágrafo {i}")
                    elif signature_start > -1 and ('PELO CONTRATANTE' in paragraph.text or 'CONTRATANTE:' in paragraph.text):
                        signature_end = i
                        logger.info(f"Fim da seção de assinaturas encontrado no parágrafo {i}")
                        break
                
                if signature_start > -1:
                    # Limpar parágrafos da seção de assinaturas antiga
                    if signature_end == -1:
                        signature_end = len(doc.paragraphs) - 5  # Aproximação se não encontrar o fim
                    
                    logger.info(f"Removendo parágrafos de assinatura de {signature_start + 1} até {signature_end - 1}")
                    
                    # Remover parágrafos antigos (em ordem reversa para não afetar índices)
                    paragraphs_to_remove = []
                    for i in range(signature_start + 1, signature_end):
                        if i < len(doc.paragraphs):
                            paragraphs_to_remove.append(doc.paragraphs[i])
                    
                    for p in reversed(paragraphs_to_remove):
                        p.clear()
                    
                    # Adicionar novas assinaturas dinamicamente
                    insert_point = signature_start + 1
                    for i, company in enumerate(empresas_bpo_lista):
                        # Adicionar linha em branco
                        new_para = doc.paragraphs[insert_point].insert_paragraph_before("")
                        
                        # Adicionar linha de assinatura (só a linha)
                        signature_line = "_" * 39
                        signature_para = doc.paragraphs[insert_point + 1].insert_paragraph_before(signature_line)
                        
                        # Adicionar nome da empresa abaixo da linha
                        name_para = doc.paragraphs[insert_point + 2].insert_paragraph_before(company['nome'])
                        
                        # Adicionar linha em branco
                        doc.paragraphs[insert_point + 3].insert_paragraph_before("")
                        
                        insert_point += 4
                        logger.info(f"Adicionada assinatura para: {company['nome']}")
                    
                    logger.info(f"Seção de assinaturas regenerada com {len(empresas_bpo_lista)} empresas")
                
            else:
                # Lógica original para contratos BPO
                # Carregar template
                doc = Document(self.template_path)
                
                # Gerar texto dinâmico das empresas BPO e lista para assinaturas
                empresas_bpo_texto_pt, empresas_bpo_texto_en, empresas_bpo_lista = self._extract_bpo_companies(company_data)
                
                # Gerar seção de assinaturas dinamicamente
                self._generate_signature_section(empresas_bpo_lista)
                
                # Recarregar template após modificar assinaturas
                doc = Document(self.template_path)
                
                # Mapeamento de campos básicos
                field_mapping = {
                    '[RAZÃO SOCIAL]': company_data['razao_social'],
                    '[CNPJ]': company_data['cnpj'],
                    '[ENDEREÇO]': company_data['endereco'],
                    '[COMPANY NAME]': company_data['razao_social'],
                    '[ADDRESS]': company_data['endereco']
                }
                
                logger.info(f"Campos a serem substituídos: {list(field_mapping.keys())}")
                
                # Substituir campos nos parágrafos (atualizado para negrito)
                total_replacements = 0
                
                # Primeiro, verificar e substituir nas tabelas (onde estão as empresas BPO)
                for table_idx, table in enumerate(doc.tables):
                    for row_idx, row in enumerate(table.rows):
                        for cell_idx, cell in enumerate(row.cells):
                            original_text = cell.text.strip()
                            new_text = original_text
                            
                            # Substituir texto em português das empresas BPO (buscar por S.JOBS)
                            if 'S.JOBS SERVIÇOS DE CONTABILIDADE LTDA.' in original_text and 'inscrita no CNPJ' in original_text:
                                logger.info(f"Substituindo célula da tabela {table_idx}[{row_idx}][{cell_idx}] com texto dinâmico das empresas BPO (PT)")
                                new_text = empresas_bpo_texto_pt
                                total_replacements += 1
                            # Substituir texto em inglês das empresas BPO (buscar por enrolled with CNPJ)
                            elif 'enrolled with CNPJ under No.' in original_text and 'CONTRACTED PARTY' in original_text:
                                logger.info(f"Substituindo célula da tabela {table_idx}[{row_idx}][{cell_idx}] com texto dinâmico das empresas BPO (EN)")
                                new_text = empresas_bpo_texto_en
                                total_replacements += 1
                            # Remover textos indesejados nas tabelas
                            elif 'ESPAÇO PROPOSITALMENTE DEIXADO EM BRANCO' in original_text:
                                new_text = original_text.replace('*******ESPAÇO PROPOSITALMENTE DEIXADO EM BRANCO *********', '')
                                new_text = new_text.replace('ESPAÇO PROPOSITALMENTE DEIXADO EM BRANCO', '')
                                total_replacements += 1
                                logger.info("Removido texto indesejado da tabela")
                            else:
                                # Substituir placeholders básicos
                                for placeholder, value in field_mapping.items():
                                    if placeholder in new_text:
                                        new_text = new_text.replace(placeholder, str(value))
                                        total_replacements += 1
                                        logger.debug(f"Substituído '{placeholder}' por '{value[:50]}...' na tabela")
                            
                            # Aplicar o novo texto se houve mudanças
                            if new_text != original_text:
                                cell.text = new_text
                
                # Depois, verificar parágrafos para outros placeholders
                for i, paragraph in enumerate(doc.paragraphs):
                    if 'ESPAÇO PROPOSITALMENTE DEIXADO EM BRANCO' in paragraph.text:
                        logger.info(f"Removendo parágrafo {i} com texto indesejado")
                        paragraph.text = ""  # Limpar o parágrafo
                        total_replacements += 1
                    elif paragraph.text:
                        for run in paragraph.runs:
                            for placeholder, value in field_mapping.items():
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))
                                    run.bold = True
                                    total_replacements += 1
                                    logger.debug(f"Substituído '{placeholder}' por '{value[:50]}...' em run do parágrafo")
                
                logger.info(f"Total de substituições realizadas: {total_replacements}")
                logger.info(f"Texto das empresas BPO (PT) gerado: {empresas_bpo_texto_pt[:400]}...")
                logger.info(f"Texto das empresas BPO (EN) gerado: {empresas_bpo_texto_en[:400]}...")
            
            # Gerar nome único para arquivo temporário
            temp_filename = f"contrato_{company_data['cnpj'].replace('.', '').replace('/', '').replace('-', '')}_{hash(company_data['razao_social']) % 10000}.docx"
            temp_filepath = os.path.join(tempfile.gettempdir(), temp_filename)
            
            # Salvar documento preenchido
            doc.save(temp_filepath)
            
            logger.info(f"Contrato gerado com sucesso: {temp_filepath}")
            return temp_filepath
            
        except Exception as e:
            logger.error(f"Erro ao gerar contrato: {str(e)}")
            raise

    def convert_to_pdf(self, docx_path: str) -> str:
        """
        Converte um arquivo .docx existente para .pdf usando LibreOffice.
        
        Args:
            docx_path: O caminho absoluto para o arquivo .docx.
            
        Returns:
            O caminho para o arquivo .pdf gerado.
        """
        try:
            # Verificar se o arquivo .docx existe
            if not os.path.exists(docx_path):
                logger.error(f"Arquivo DOCX não encontrado: {docx_path}")
                raise FileNotFoundError(f"Arquivo DOCX não encontrado: {docx_path}")
            
            logger.info(f"Arquivo DOCX encontrado: {docx_path}, tamanho: {os.path.getsize(docx_path)} bytes")
            
            # Definir caminho de saída do PDF
            pdf_path = docx_path.replace('.docx', '.pdf')
            logger.info(f"Convertendo {docx_path} para PDF em {pdf_path}...")
            
            # Usar LibreOffice em modo headless para converter para PDF
            import subprocess
            import shutil
            
            # Comando padrão para Linux (onde o Docker está rodando)
            soffice_cmd = "soffice"
            
            # Verificar se o LibreOffice está disponível
            if shutil.which(soffice_cmd):
                logger.info(f"LibreOffice encontrado, usando para conversão")
                
                # Comando para converter DOCX para PDF
                cmd = [
                    soffice_cmd,
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(pdf_path),
                    docx_path
                ]
                
                # Executar o comando com timeout para evitar bloqueios
                logger.info(f"Executando comando: {' '.join(cmd)}")
                process = subprocess.Popen(
                    cmd, 
                    stdout=subprocess.PIPE, 
                    stderr=subprocess.PIPE
                )
                
                stdout, stderr = process.communicate(timeout=30)
                
                # Verificar o resultado da conversão
                if process.returncode != 0:
                    logger.error(f"Erro ao converter com LibreOffice: {stderr.decode()}")
                    raise RuntimeError(f"Falha na conversão para PDF: {stderr.decode()}")
                
                logger.info(f"Saída da conversão: {stdout.decode()}")
                
                # Verificar se o PDF foi gerado corretamente
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                    logger.info(f"PDF gerado com sucesso: {pdf_path}, tamanho: {os.path.getsize(pdf_path)} bytes")
                    return pdf_path
                else:
                    logger.error(f"PDF não foi gerado ou está vazio: {pdf_path}")
                    raise RuntimeError(f"PDF não foi gerado corretamente: {pdf_path}")
            else:
                # Se LibreOffice não estiver disponível, retornar o arquivo DOCX como alternativa
                logger.warning("LibreOffice não está disponível, retornando arquivo DOCX original")
                return docx_path
            
        except subprocess.TimeoutExpired as timeout_error:
            if 'process' in locals():
                process.kill()
            logger.error("Timeout ao executar LibreOffice")
            return docx_path
            
        except Exception as e:
            logger.error(f"Erro ao converter para PDF: {str(e)}")
            # Em caso de erro, retorna o arquivo DOCX original
            return docx_path
    
    def get_template_fields(self) -> list:
        """
        Extrai os campos do template que precisam ser preenchidos
        
        Returns:
            Lista de placeholders encontrados no template
        """
        try:
            doc = Document(self.template_path)
            placeholders = set()
            
            # Buscar em parágrafos
            for paragraph in doc.paragraphs:
                text = paragraph.text
                # Buscar padrão [CAMPO]
                import re
                matches = re.findall(r'\[([^\]]+)\]', text)
                placeholders.update(matches)
            
            # Buscar em tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text
                        matches = re.findall(r'\[([^\]]+)\]', text)
                        placeholders.update(matches)
            
            return sorted(list(placeholders))
            
        except Exception as e:
            logger.error(f"Erro ao extrair campos do template: {str(e)}")
            return []
    
    def extract_text_content(self, file_path: str) -> str:
        """
        Extrai o conteúdo textual completo de um documento Word na ordem original
        
        Args:
            file_path: Caminho para o arquivo .docx
            
        Returns:
            Conteúdo textual completo do documento com formatação preservada
        """
        try:
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.table import Table
            from docx.text.paragraph import Paragraph
            import json
            
            doc = Document(file_path)
            content_lines = []
            
            logger.info("Extraindo conteúdo do documento para edição...")
            
            # Processar elementos na ordem que aparecem no documento
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # É um parágrafo
                    paragraph = Paragraph(element, doc)
                    paragraph_text = paragraph.text.strip()
                    
                    # Incluir TODOS os parágrafos, mesmo vazios (para manter correspondência)
                    if paragraph_text:
                        content_lines.append(paragraph_text)
                        logger.debug(f"Parágrafo extraído: '{paragraph_text}'")
                    else:
                        # Parágrafos vazios também são importantes para manter estrutura
                        content_lines.append("")
                    
                elif isinstance(element, CT_Tbl):
                    # É uma tabela - converter para JSON
                    table = Table(element, doc)
                    
                    # Verificar se a tabela tem conteúdo
                    has_content = False
                    table_rows = []
                    
                    for row in table.rows:
                        row_cells = []
                        for cell in row.cells:
                            # Preservar quebras de linha dentro das células
                            cell_text = cell.text.strip()
                            if cell_text:
                                has_content = True
                            row_cells.append(cell_text)
                        table_rows.append(row_cells)
                    
                    if has_content and table_rows:
                        # Converter tabela para JSON que o frontend pode processar
                        table_json = json.dumps(table_rows, ensure_ascii=False)
                        logger.debug(f"Tabela extraída com {len(table_rows)} linhas e {max(len(row) for row in table_rows) if table_rows else 0} colunas")
                        
                        # Adicionar quebras de linha antes e depois para garantir que fique isolado
                        content_lines.append('')  # Linha em branco antes
                        content_lines.append(f"[TABELA_JSON]{table_json}[/TABELA_JSON]")
                        content_lines.append('')  # Linha em branco depois
            
            final_content = '\n'.join(content_lines)
            logger.info(f"Extração concluída: {len(content_lines)} linhas totais")
            
            return final_content
            
        except Exception as e:
            logger.error(f"Erro ao extrair conteúdo: {str(e)}")
            # Fallback simples
            try:
                doc = Document(file_path)
                content_lines = []
                for paragraph in doc.paragraphs:
                    content_lines.append(paragraph.text)
                return '\n'.join(content_lines)
            except:
                return ""
    
    def apply_text_edits(self, contract_path: str, new_content: str) -> str:
        """
        Aplica edições de texto completo ao contrato, reconstruindo tabelas corretamente
        
        Args:
            contract_path: Caminho do contrato original
            new_content: Novo conteúdo completo do contrato
            
        Returns:
            Caminho do novo arquivo editado
        """
        try:
            logger.info(f"Aplicando edições de texto completo ao contrato: {contract_path}")
            
            # Criar novo documento
            doc = Document()
            
            # Dividir conteúdo em linhas
            lines = new_content.split('\n')
            
            i = 0
            while i < len(lines):
                line = lines[i]
                
                # Verificar se é início de uma tabela (formato JSON)
                if "[TABELA_JSON]" in line:
                    import json
                    import re
                    
                    # Extrair JSON da linha
                    json_match = re.search(r'\[TABELA_JSON\](.*?)\[/TABELA_JSON\]', line)
                    if json_match:
                        try:
                            table_data = json.loads(json_match.group(1))
                            
                            if table_data and len(table_data) > 0:
                                # Determinar número de colunas
                                max_cols = max(len(row) for row in table_data) if table_data else 0
                                
                                # Criar tabela
                                table = doc.add_table(rows=len(table_data), cols=max_cols)
                                table.style = 'Table Grid'
                                
                                # Preencher dados
                                for row_idx, row_data in enumerate(table_data):
                                    for col_idx, cell_data in enumerate(row_data):
                                        if col_idx < max_cols and row_idx < len(table.rows):
                                            table.rows[row_idx].cells[col_idx].text = str(cell_data)
                                
                                logger.info(f"Tabela JSON reconstruída com {len(table_data)} linhas e {max_cols} colunas")
                        except json.JSONDecodeError as e:
                            logger.error(f"Erro ao decodificar JSON da tabela: {e}")
                            # Em caso de erro, adicionar como texto normal
                            doc.add_paragraph(line)
                    i += 1
                
                # Verificar se é início de uma tabela (novo formato)
                if line.strip() == "┌─ TABELA ─┐":
                    # Coletar linhas da tabela até encontrar o fechamento
                    table_lines = []
                    i += 1
                    
                    # Coletar todas as linhas da tabela
                    while i < len(lines):
                        current_line = lines[i]
                        if current_line.strip().startswith('└─') and '─┘' in current_line:
                            # Fim da tabela
                            break
                        elif current_line.strip().startswith('│') and current_line.strip().endswith('│'):
                            # Linha de dados da tabela
                            table_lines.append(current_line)
                        elif current_line.strip().startswith('├─') and '─┤' in current_line:
                            # Linha separadora - ignorar
                            pass
                        i += 1
                    
                    # Processar dados da tabela
                    if table_lines:
                        table_data = []
                        for table_line in table_lines:
                            # Remover │ do início e fim, dividir por │
                            clean_line = table_line.strip()
                            if clean_line.startswith('│') and clean_line.endswith('│'):
                                clean_line = clean_line[1:-1]  # Remover │ do início e fim
                                cells = [cell.strip() for cell in clean_line.split('│')]
                                if cells and any(cell.strip() for cell in cells):  # Se tem conteúdo
                                    table_data.append(cells)
                        
                        # Criar tabela se tiver dados
                        if table_data:
                            # Determinar número máximo de colunas
                            max_cols = max(len(row) for row in table_data)
                            
                            # Criar tabela
                            table = doc.add_table(rows=len(table_data), cols=max_cols)
                            
                            # Preencher dados
                            for row_idx, row_data in enumerate(table_data):
                                for col_idx, cell_data in enumerate(row_data):
                                    if col_idx < max_cols:
                                        table.cell(row_idx, col_idx).text = cell_data.strip()
                            
                            # Adicionar parágrafo vazio após tabela
                            doc.add_paragraph()
                            
                            logger.info(f"Tabela reconstruída com {len(table_data)} linhas e {max_cols} colunas")
                    
                    i += 1  # Pular linha de fechamento da tabela
                    
                # Verificar se é início de uma tabela (formato antigo com [INÍCIO_TABELA])
                elif line.strip() == "[INÍCIO_TABELA]":
                    # Coletar linhas da tabela até encontrar [FIM_TABELA]
                    table_rows = []
                    i += 1
                    
                    while i < len(lines) and lines[i].strip() != "[FIM_TABELA]":
                        if lines[i].strip():  # Ignorar linhas vazias dentro da tabela
                            # Dividir por [CÉLULA] para obter células
                            cells = lines[i].split(" [CÉLULA] ")
                            if cells:
                                table_rows.append(cells)
                        i += 1
                    
                    # Criar tabela se tiver dados
                    if table_rows:
                        # Determinar número máximo de colunas
                        max_cols = max(len(row) for row in table_rows)
                        
                        # Criar tabela
                        table = doc.add_table(rows=len(table_rows), cols=max_cols)
                        
                        # Preencher dados
                        for row_idx, row_data in enumerate(table_rows):
                            for col_idx, cell_data in enumerate(row_data):
                                if col_idx < max_cols:
                                    table.cell(row_idx, col_idx).text = cell_data.strip()
                        
                        # Adicionar parágrafo vazio após tabela
                        doc.add_paragraph()
                        
                        logger.info(f"Tabela (formato antigo) reconstruída com {len(table_rows)} linhas e {max_cols} colunas")
                    
                    i += 1  # Pular [FIM_TABELA]
                    
                else:
                    # Linha normal - adicionar como parágrafo
                    if line.strip():
                        doc.add_paragraph(line)
                    else:
                        doc.add_paragraph()  # Parágrafo vazio para espaçamento
                    i += 1
            
            # Salvar arquivo editado
            base_filename = os.path.basename(contract_path)
            name_without_ext = os.path.splitext(base_filename)[0]
            edited_filename = f"{name_without_ext}_editado.docx"
            edited_filepath = os.path.join(tempfile.gettempdir(), edited_filename)
            
            doc.save(edited_filepath)
            logger.info(f"Contrato editado salvo: {edited_filepath}")
            
            return edited_filepath
            
        except Exception as e:
            logger.error(f"Erro ao aplicar edições de texto: {str(e)}")
            raise

    def _apply_xml_edits(self, contract_path: str, new_content: str, output_path: str) -> str:
        """
        Aplica edições usando manipulação XML AVANÇADA com preservação total de formatação
        """
        try:
            import zipfile
            from lxml import etree
            import tempfile
            import json
            import re
            import shutil
            from copy import deepcopy
            
            # Word é um ZIP com XML dentro - vamos manipular diretamente
            temp_dir = tempfile.mkdtemp()
            
            try:
                # Extrair o ZIP do Word
                with zipfile.ZipFile(output_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)
                
                logger.info("📂 Documento Word extraído como XML")
                
                # Manipular o document.xml (conteúdo principal)
                doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
                
                # Ler o XML original
                with open(doc_xml_path, 'r', encoding='utf-8') as f:
                    xml_content = f.read()
                
                # Parse do XML
                root = etree.fromstring(xml_content.encode('utf-8'))
                
                # Namespace do Word
                ns = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
                    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
                    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006'
                }
                
                logger.info("🔍 XML parseado com sucesso")
                
                # Encontrar o body
                body = root.find('.//w:body', ns)
                if body is None:
                    raise Exception("Body não encontrado no XML")
                
                # NOVA ESTRATÉGIA: Mapear conteúdo original para novos valores
                original_content = self.extract_text_content(contract_path)
                original_lines = [line.strip() for line in original_content.split('\n') if line.strip()]
                new_lines = [line.strip() for line in new_content.split('\n') if line.strip()]
                
                logger.info(f"📋 Mapeando {len(original_lines)} → {len(new_lines)} linhas")
                
                # Separar linhas normais de tabelas JSON
                normal_mapping = {}
                table_mapping = {}
                changes_count = 0
                
                for i, orig_line in enumerate(original_lines):
                    if i < len(new_lines):
                        if orig_line != new_lines[i]:
                            if "[TABELA_JSON]" in orig_line:
                                table_mapping[orig_line] = new_lines[i]
                                logger.info(f"📊 Mapeamento tabela {len(table_mapping)}: '{orig_line[:50]}...' → '{new_lines[i][:50]}...'")
                            else:
                                normal_mapping[orig_line] = new_lines[i]
                                logger.info(f"📝 Mapeamento texto {len(normal_mapping)}: '{orig_line[:30]}...' → '{new_lines[i][:30]}...'")
                            changes_count += 1
                    else:
                        if "[TABELA_JSON]" in orig_line:
                            table_mapping[orig_line] = ""  # Tabela removida
                        else:
                            normal_mapping[orig_line] = ""  # Linha removida
                        changes_count += 1
                
                logger.info(f"📊 Total: {len(normal_mapping)} textos + {len(table_mapping)} tabelas = {changes_count} mudanças")
                
                # Adicionar linhas novas que não existiam
                extra_lines = []
                if len(new_lines) > len(original_lines):
                    extra_lines = new_lines[len(original_lines):]
                    logger.info(f"➕ {len(extra_lines)} linhas adicionais detectadas")
                
                # Processar textos normais
                if normal_mapping:
                    logger.info(f"� Processando {len(normal_mapping)} mudanças de texto...")
                    self._update_xml_content_preserving_format(body, normal_mapping, ns)
                
                # Processar tabelas JSON separadamente
                if table_mapping:
                    logger.info(f"📊 Processando {len(table_mapping)} mudanças de tabela...")
                    self._update_xml_tables(body, table_mapping, ns)
                
                # Adicionar linhas extras se houver
                if extra_lines:
                    sect_pr = body.find('.//w:sectPr', ns)
                    
                    for extra_line in extra_lines:
                        if extra_line.strip():
                            # Criar parágrafo simples para linha extra
                            para_xml = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                                <w:r>
                                    <w:t>{extra_line}</w:t>
                                </w:r>
                            </w:p>'''
                            new_para = etree.fromstring(para_xml)
                            if sect_pr is not None:
                                body.insert(-1, new_para)
                            else:
                                body.append(new_para)
                    logger.info(f"➕ {len(extra_lines)} linhas extras adicionadas")
                
                # Verificar quantos mapeamentos sobraram
                remaining_normal = len(normal_mapping) if normal_mapping else 0
                remaining_tables = len(table_mapping) if table_mapping else 0
                
                if remaining_normal > 0:
                    logger.warning(f"⚠️ {remaining_normal} mapeamentos de texto não aplicados:")
                    for i, (orig, new) in enumerate(list(normal_mapping.items())[:3]):
                        logger.warning(f"  Texto {i+1}. '{orig[:40]}...' → '{new[:40]}...'")
                
                if remaining_tables > 0:
                    logger.warning(f"⚠️ {remaining_tables} mapeamentos de tabela não aplicados:")
                    for i, (orig, new) in enumerate(list(table_mapping.items())[:3]):
                        logger.warning(f"  Tabela {i+1}. '{orig[:40]}...' → '{new[:40]}...'")
                
                if remaining_normal == 0 and remaining_tables == 0:
                    logger.info("✅ Todos os mapeamentos foram aplicados com sucesso")
                
                # Salvar XML modificado
                modified_xml = etree.tostring(root, encoding='utf-8', xml_declaration=True, pretty_print=True)
                with open(doc_xml_path, 'wb') as f:
                    f.write(modified_xml)
                
                # Recriar o ZIP/DOCX
                with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
                    for root_dir, dirs, files in os.walk(temp_dir):
                        for file in files:
                            file_path = os.path.join(root_dir, file)
                            arc_name = os.path.relpath(file_path, temp_dir)
                            zip_ref.write(file_path, arc_name)
                
                logger.info(f"🎉 EDIÇÕES XML AVANÇADAS APLICADAS: {output_path}")
                return output_path
                
            finally:
                # Limpar diretório temporário
                shutil.rmtree(temp_dir, ignore_errors=True)
                
        except Exception as e:
            logger.error(f"❌ ERRO na aplicação XML avançada: {str(e)}")
            raise

    def _update_xml_tables(self, body, table_mapping, ns):
        """
        Atualiza especificamente tabelas JSON no XML
        """
        try:
            import json
            import re
            
            # Encontrar todas as tabelas no documento
            tables = body.findall('.//w:tbl', ns)
            logger.info(f"📊 Encontradas {len(tables)} tabelas no documento")
            
            # Para cada mapeamento de tabela
            for orig_table_json, new_table_json in table_mapping.items():
                # Extrair dados JSON da string original e nova
                orig_match = re.search(r'\[TABELA_JSON\](.*?)\[/TABELA_JSON\]', orig_table_json)
                new_match = re.search(r'\[TABELA_JSON\](.*?)\[/TABELA_JSON\]', new_table_json)
                
                if orig_match and new_match:
                    try:
                        orig_data = json.loads(orig_match.group(1))
                        new_data = json.loads(new_match.group(1))
                        
                        logger.info(f"📊 Atualizando tabela: {len(orig_data)}x{len(orig_data[0]) if orig_data else 0} → {len(new_data)}x{len(new_data[0]) if new_data else 0}")
                        
                        # Procurar tabela correspondente no XML
                        for table in tables:
                            rows = table.findall('.//w:tr', ns)
                            
                            # Verificar se a tabela corresponde aos dados originais
                            if len(rows) >= len(orig_data):
                                matches = 0
                                for row_idx, orig_row in enumerate(orig_data):
                                    if row_idx < len(rows):
                                        cells = rows[row_idx].findall('.//w:tc', ns)
                                        if len(cells) >= len(orig_row):
                                            for col_idx, orig_cell in enumerate(orig_row):
                                                if col_idx < len(cells):
                                                    cell_text = ''.join(t.text or '' for t in cells[col_idx].findall('.//w:t', ns))
                                                    if str(orig_cell).strip() == cell_text.strip():
                                                        matches += 1
                                
                                # Se a maioria das células corresponde, essa é nossa tabela
                                total_cells = sum(len(row) for row in orig_data)
                                if matches >= total_cells * 0.7:  # 70% de correspondência
                                    logger.info(f"✅ Tabela encontrada! {matches}/{total_cells} células correspondem")
                                    
                                    # Atualizar com novos dados
                                    for row_idx, new_row in enumerate(new_data):
                                        if row_idx < len(rows):
                                            cells = rows[row_idx].findall('.//w:tc', ns)
                                            for col_idx, new_cell in enumerate(new_row):
                                                if col_idx < len(cells):
                                                    # Atualizar texto da célula preservando formatação
                                                    text_elements = cells[col_idx].findall('.//w:t', ns)
                                                    if text_elements:
                                                        # Limpar elementos extras
                                                        for t_elem in text_elements[1:]:
                                                            t_elem.text = ""
                                                        # Aplicar novo texto
                                                        text_elements[0].text = str(new_cell)
                                    
                                    logger.info(f"🎉 Tabela atualizada com sucesso!")
                                    break
                        
                    except json.JSONDecodeError as e:
                        logger.error(f"Erro ao processar JSON da tabela: {e}")
                
        except Exception as e:
            logger.error(f"Erro ao atualizar tabelas XML: {e}")

    def _update_xml_content_preserving_format(self, element, content_mapping, ns):
        """
        Atualiza conteúdo XML preservando totalmente a formatação original
        """
        try:
            # Processar parágrafos
            if element.tag.endswith('}p'):
                # Extrair texto atual do parágrafo
                current_text = ''.join(t.text or '' for t in element.findall('.//w:t', ns))
                current_text = current_text.strip()
                
                # DEBUG: Log do texto encontrado
                if current_text and len(current_text) > 10:  # Apenas textos significativos
                    logger.debug(f"🔍 Processando parágrafo: '{current_text[:50]}...'")
                
                # Verificar se há mapeamento EXATO para este texto
                if current_text in content_mapping:
                    new_text = content_mapping[current_text]
                    if new_text != current_text:
                        # Preservar TODA a formatação, apenas trocar o texto
                        text_elements = element.findall('.//w:t', ns)
                        if text_elements:
                            # Limpar todos os elementos de texto exceto o primeiro
                            for t_elem in text_elements[1:]:
                                t_elem.text = ""
                            # Colocar todo o novo texto no primeiro elemento
                            text_elements[0].text = new_text
                            logger.info(f"✅ Texto atualizado: '{current_text[:30]}...' → '{new_text[:30]}...'")
                    # Remover do mapeamento para não reutilizar
                    del content_mapping[current_text]
                else:
                    # Se não encontrar correspondência exata, tentar busca parcial
                    for original_text, new_text in list(content_mapping.items()):
                        # Ignorar tabelas JSON
                        if "[TABELA_JSON]" in original_text or "[TABELA_JSON]" in current_text:
                            continue
                            
                        # Busca parcial: se o texto original contém o texto atual ou vice-versa
                        if (current_text in original_text or original_text in current_text) and len(current_text) > 5:
                            text_elements = element.findall('.//w:t', ns)
                            if text_elements:
                                # Substituir texto
                                for t_elem in text_elements[1:]:
                                    t_elem.text = ""
                                text_elements[0].text = new_text
                                logger.info(f"🔄 Texto atualizado (parcial): '{current_text[:30]}...' → '{new_text[:30]}...'")
                                del content_mapping[original_text]
                                break
            
            # Processar tabelas
            elif element.tag.endswith('}tbl'):
                # Para tabelas, processar cada célula individualmente
                cells = element.findall('.//w:tc', ns)
                for cell in cells:
                    self._update_xml_content_preserving_format(cell, content_mapping, ns)
            
            # Recursivamente processar elementos filhos
            for child in element:
                self._update_xml_content_preserving_format(child, content_mapping, ns)
                
        except Exception as e:
            logger.error(f"Erro ao atualizar elemento XML: {e}")

    def apply_selective_edits(self, contract_path: str, new_content: str) -> str:
        """
        Aplica edições de forma inteligente: só modifica se realmente mudou algo.
        Estratégia: compara conteúdo e só reconstrói se necessário.
        
        Args:
            contract_path: Caminho do contrato original
            new_content: Novo conteúdo completo editado
            
        Returns:
            Caminho do novo arquivo editado
        """
        try:
            logger.info(f"🔧 INICIANDO aplicação inteligente: {contract_path}")
            
            import shutil
            import tempfile
            import json
            import re
            
            # Criar cópia do arquivo original
            base_filename = os.path.basename(contract_path)
            name_without_ext = os.path.splitext(base_filename)[0]
            edited_filename = f"{name_without_ext}_editado.docx"
            edited_filepath = os.path.join(tempfile.gettempdir(), edited_filename)
            
            # SEMPRE copiar o original primeiro
            shutil.copy2(contract_path, edited_filepath)
            logger.info(f"📋 Arquivo copiado: {edited_filepath}")
            
            # 1. EXTRAIR CONTEÚDO ATUAL DO DOCUMENTO
            current_content = self.extract_text_content(contract_path)
            logger.info(f"📄 Conteúdo atual extraído: {len(current_content)} chars")
            
            # 2. COMPARAR SE REALMENTE MUDOU ALGO
            current_lines = [line.strip() for line in current_content.split('\n')]
            new_lines = [line.strip() for line in new_content.split('\n')]
            
            # Log para debug
            logger.info(f"🔍 DEBUG COMPARAÇÃO:")
            logger.info(f"  - Linhas originais: {len(current_lines)}")
            logger.info(f"  - Linhas novas: {len(new_lines)}")
            logger.info(f"  - Primeiras 3 linhas originais: {current_lines[:3]}")
            logger.info(f"  - Primeiras 3 linhas novas: {new_lines[:3]}")
            
            # Remover linhas vazias para comparação
            current_clean = [line for line in current_lines if line]
            new_clean = [line for line in new_lines if line]
            
            logger.info(f"🔍 DEBUG PÓS-LIMPEZA:")
            logger.info(f"  - Linhas originais limpas: {len(current_clean)}")
            logger.info(f"  - Linhas novas limpas: {len(new_clean)}")
            
            # Verificar se houve mudanças significativas
            content_changed = False
            
            if len(current_clean) != len(new_clean):
                content_changed = True
                logger.info(f"🔄 Mudança detectada: número de linhas {len(current_clean)} → {len(new_clean)}")
            else:
                # Comparar linha por linha (ignorando formatações JSON de tabela)
                changes_found = []
                for i, (current_line, new_line) in enumerate(zip(current_clean, new_clean)):
                    # Pular comparação de linhas JSON de tabela (são geradas automaticamente)
                    if "[TABELA_JSON]" in current_line or "[TABELA_JSON]" in new_line:
                        continue
                    
                    if current_line != new_line:
                        content_changed = True
                        changes_found.append({
                            'line': i,
                            'before': current_line[:100],
                            'after': new_line[:100]
                        })
                        if len(changes_found) <= 5:  # Log apenas as primeiras 5 diferenças
                            logger.info(f"🔄 Mudança detectada na linha {i}:")
                            logger.info(f"  ANTES: '{current_line[:100]}...'")
                            logger.info(f"  DEPOIS: '{new_line[:100]}...'")
                
                if changes_found:
                    logger.info(f"🔄 Total de mudanças detectadas: {len(changes_found)}")
                else:
                    logger.info("🔍 Nenhuma diferença encontrada na comparação linha por linha")
            
            # 3. SE NÃO HOUVE MUDANÇAS, RETORNAR CÓPIA SIMPLES
            if not content_changed:
                # VERIFICAÇÃO ADICIONAL: Comparar tamanho bruto do conteúdo
                content_size_diff = abs(len(current_content) - len(new_content))
                logger.info(f"🔍 Diferença de tamanho: {content_size_diff} caracteres")
                
                if content_size_diff > 50:  # Se a diferença for significativa
                    logger.info("🔄 DIFERENÇA DE TAMANHO SIGNIFICATIVA - Aplicando edições mesmo sem mudanças detectadas linha por linha")
                    content_changed = True
                else:
                    logger.info("✅ NENHUMA MUDANÇA DETECTADA - Retornando arquivo original preservado")
                    return edited_filepath
            
            # 4. SE HOUVE MUDANÇAS, APLICAR ESTRATÉGIA XML
            logger.info("🚀 MUDANÇAS DETECTADAS - Aplicando edições com XML...")
            
            return self._apply_xml_edits(contract_path, new_content, edited_filepath)
            
        except Exception as e:
            logger.error(f"❌ ERRO na aplicação inteligente: {str(e)}")
            logger.exception("Detalhes completos:")
            
            # Fallback: retornar cópia simples do original
            import shutil
            base_filename = os.path.basename(contract_path)
            name_without_ext = os.path.splitext(base_filename)[0]
            fallback_filename = f"{name_without_ext}_editado.docx"
            fallback_filepath = os.path.join(tempfile.gettempdir(), fallback_filename)
            shutil.copy2(contract_path, fallback_filepath)
            logger.info(f"📋 Fallback: retornando cópia preservada do original")
            return fallback_filepath

    def apply_selective_edits_OLD(self, contract_path: str, new_content: str) -> str:
        """
        ABORDAGEM FINAL: Manipulação direta do XML do documento Word.
        Esta é a única forma 100% confiável de preservar formatação E aplicar edições.
        
        Args:
            contract_path: Caminho do contrato original
            new_content: Novo conteúdo completo editado
            
        Returns:
            Caminho do novo arquivo editado
        """
        try:
            logger.info(f"� ABORDAGEM XML DIRETA: {contract_path}")
            
            import shutil
            import tempfile
            import json
            import re
            import zipfile
            from lxml import etree
            
            # Criar cópia do arquivo original
            base_filename = os.path.basename(contract_path)
            name_without_ext = os.path.splitext(base_filename)[0]
            edited_filename = f"{name_without_ext}_editado.docx"
            edited_filepath = os.path.join(tempfile.gettempdir(), edited_filename)
            
            shutil.copy2(contract_path, edited_filepath)
            logger.info(f"📋 Arquivo copiado: {edited_filepath}")
            
            # Word é um ZIP com XML dentro - vamos manipular diretamente
            temp_dir = tempfile.mkdtemp()
            
            try:
                # Extrair o ZIP do Word
                with zipfile.ZipFile(edited_filepath, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)
                
                logger.info("📂 Documento Word extraído como XML")
                
                # Manipular o document.xml (conteúdo principal)
                doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
                
                # Ler o XML original
                with open(doc_xml_path, 'r', encoding='utf-8') as f:
                    xml_content = f.read()
                
                # Parse do XML
                root = etree.fromstring(xml_content.encode('utf-8'))
                
                # Namespace do Word
                ns = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
                }
                
                logger.info("🔍 XML parseado com sucesso")
                
                # Encontrar o body
                body = root.find('.//w:body', ns)
                if body is None:
                    raise Exception("Body não encontrado no XML")
                
                # Coletar TODOS os elementos formatados (parágrafos e tabelas)
                formatted_elements = []
                
                for elem in body:
                    if elem.tag.endswith('}p'):  # Parágrafo
                        formatted_elements.append(('paragraph', elem))
                    elif elem.tag.endswith('}tbl'):  # Tabela
                        formatted_elements.append(('table', elem))
                
                logger.info(f"📋 Coletados {len(formatted_elements)} elementos formatados")
                
                # Processar novo conteúdo
                new_lines = new_content.split('\n')
                logger.info(f"📝 Processando {len(new_lines)} linhas do novo conteúdo")
                
                # Limpar body mantendo apenas sectPr (configurações de seção)
                sect_pr = body.find('.//w:sectPr', ns)
                body.clear()
                if sect_pr is not None:
                    body.append(sect_pr)
                
                # Aplicar novo conteúdo usando elementos formatados como template
                para_templates = [elem[1] for elem in formatted_elements if elem[0] == 'paragraph']
                table_templates = [elem[1] for elem in formatted_elements if elem[0] == 'table']
                
                para_template_index = 0
                table_template_index = 0
                
                i = 0
                while i < len(new_lines):
                    line = new_lines[i]
                    
                    if "[TABELA_JSON]" in line:
                        # Processar tabela
                        json_match = re.search(r'\[TABELA_JSON\](.*?)\[/TABELA_JSON\]', line)
                        if json_match:
                            try:
                                table_data = json.loads(json_match.group(1))
                                
                                if table_data and table_template_index < len(table_templates):
                                    # Usar template de tabela disponível
                                    table_template = table_templates[table_template_index]
                                    table_template_index += 1
                                    
                                    # Clonar template da tabela
                                    new_table = etree.fromstring(etree.tostring(table_template))
                                    
                                    # Limpar dados antigos, mas manter estrutura
                                    rows = new_table.findall('.//w:tr', ns)
                                    
                                    # Ajustar número de linhas se necessário
                                    current_rows = len(rows)
                                    needed_rows = len(table_data)
                                    
                                    if needed_rows > current_rows:
                                        # Adicionar linhas clonando a última
                                        if rows:
                                            last_row = rows[-1]
                                            for _ in range(needed_rows - current_rows):
                                                new_row = etree.fromstring(etree.tostring(last_row))
                                                # Encontrar o elemento tbl e adicionar a nova linha
                                                tbl = new_table if new_table.tag.endswith('}tbl') else new_table.find('.//w:tbl', ns)
                                                if tbl is not None:
                                                    tbl.append(new_row)
                                                    rows.append(new_row)
                                    
                                    # Aplicar novos dados mantendo formatação
                                    for row_idx, row_data in enumerate(table_data):
                                        if row_idx < len(rows):
                                            cells = rows[row_idx].findall('.//w:tc', ns)
                                            for col_idx, cell_data in enumerate(row_data):
                                                if col_idx < len(cells):
                                                    # Encontrar todos os elementos de texto na célula
                                                    text_elements = cells[col_idx].findall('.//w:t', ns)
                                                    if text_elements:
                                                        # Limpar texto de todos os elementos
                                                        for t_elem in text_elements[1:]:
                                                            t_elem.text = ""
                                                        # Colocar novo texto no primeiro elemento
                                                        text_elements[0].text = str(cell_data)
                                                    else:
                                                        # Se não há elementos de texto, criar um básico
                                                        tc = cells[col_idx]
                                                        # Encontrar ou criar um parágrafo
                                                        p = tc.find('.//w:p', ns)
                                                        if p is not None:
                                                            # Limpar parágrafo
                                                            for child in list(p):
                                                                p.remove(child)
                                                            # Criar run com texto
                                                            r_elem = etree.SubElement(p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                                                            t_elem = etree.SubElement(r_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                                                            t_elem.text = str(cell_data)
                                    
                                    body.insert(-1 if sect_pr is not None else len(body), new_table)
                                    logger.info(f"📊 Tabela XML aplicada na ordem: {len(table_data)} linhas")
                                
                            except Exception as e:
                                logger.error(f"Erro ao processar tabela: {e}")
                                # Em caso de erro, criar parágrafo simples
                                if para_template_index < len(para_templates):
                                    para_template = para_templates[para_template_index]
                                    para_template_index += 1
                                    new_para = etree.fromstring(etree.tostring(para_template))
                                    # Limpar e colocar linha como texto
                                    text_elements = new_para.findall('.//w:t', ns)
                                    if text_elements:
                                        for t_elem in text_elements:
                                            t_elem.text = ""
                                        text_elements[0].text = line
                                    body.insert(-1 if sect_pr is not None else len(body), new_para)
                        
                        i += 1
                    
                    else:
                        # Processar parágrafo normal
                        if para_template_index < len(para_templates):
                            para_template = para_templates[para_template_index]
                            para_template_index += 1
                            
                            # Clonar template do parágrafo
                            new_para = etree.fromstring(etree.tostring(para_template))
                            
                            # Limpar texto antigo, mas manter formatação
                            text_elements = new_para.findall('.//w:t', ns)
                            if text_elements:
                                # Limpar todos os textos exceto o primeiro
                                for t_elem in text_elements[1:]:
                                    t_elem.text = ""
                                # Colocar novo texto no primeiro elemento
                                text_elements[0].text = line if line.strip() else ""
                            else:
                                # Se não há elementos de texto, criar estrutura básica
                                if line.strip():
                                    # Encontrar ou criar run
                                    r = new_para.find('.//w:r', ns)
                                    if r is None:
                                        r = etree.SubElement(new_para, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                                    # Criar elemento de texto
                                    t = etree.SubElement(r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                                    t.text = line
                            
                            body.insert(-1 if sect_pr is not None else len(body), new_para)
                            logger.debug(f"📝 Parágrafo XML na ordem: '{line[:30]}...'")
                        else:
                            # Sem mais templates, criar parágrafo básico
                            if line.strip():
                                para_xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:t>{line}</w:t></w:r></w:p>'
                                new_para = etree.fromstring(para_xml)
                                body.insert(-1 if sect_pr is not None else len(body), new_para)
                            else:
                                # Parágrafo vazio
                                para_xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"></w:p>'
                                new_para = etree.fromstring(para_xml)
                                body.insert(-1 if sect_pr is not None else len(body), new_para)
                        
                        i += 1
                
                # Salvar XML modificado
                modified_xml = etree.tostring(root, encoding='utf-8', xml_declaration=True, pretty_print=True)
                with open(doc_xml_path, 'wb') as f:
                    f.write(modified_xml)
                
                logger.info("💾 XML modificado salvo")
                
                # Recriar o ZIP/DOCX
                with zipfile.ZipFile(edited_filepath, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
                    for root_dir, dirs, files in os.walk(temp_dir):
                        for file in files:
                            file_path = os.path.join(root_dir, file)
                            arc_name = os.path.relpath(file_path, temp_dir)
                            zip_ref.write(file_path, arc_name)
                
                logger.info(f"🎉 DOCUMENTO FINAL CRIADO: {edited_filepath}")
                
                return edited_filepath
                
            finally:
                # Limpar diretório temporário
                shutil.rmtree(temp_dir, ignore_errors=True)
            
        except Exception as e:
            logger.error(f"❌ ERRO na manipulação XML: {str(e)}")
            logger.exception("Detalhes completos:")
            
            # Fallback final: tentar método anterior
            try:
                return self._fallback_simple_replacement(contract_path, new_content)
            except:
                # Último recurso: retornar original
                import shutil
                base_filename = os.path.basename(contract_path)
                name_without_ext = os.path.splitext(base_filename)[0]
                fallback_filename = f"{name_without_ext}_editado.docx"
                fallback_filepath = os.path.join(tempfile.gettempdir(), fallback_filename)
                shutil.copy2(contract_path, fallback_filepath)
                return fallback_filepath
    
    def _fallback_simple_replacement(self, contract_path: str, new_content: str) -> str:
        """Método fallback simples"""
        import shutil
        import tempfile
        
        base_filename = os.path.basename(contract_path)
        name_without_ext = os.path.splitext(base_filename)[0]
        edited_filename = f"{name_without_ext}_editado.docx"
        edited_filepath = os.path.join(tempfile.gettempdir(), edited_filename)
        
        shutil.copy2(contract_path, edited_filepath)
        
        # Método simples: só aplicar novo conteúdo
        doc = Document(edited_filepath)
        
        # Limpar conteúdo
        body = doc.element.body
        for element in list(body):
            body.remove(element)
        
        # Adicionar novo conteúdo linha por linha
        for line in new_content.split('\n'):
            if "[TABELA_JSON]" in line:
                continue  # Pular tabelas no fallback
            doc.add_paragraph(line)
        
        doc.save(edited_filepath)
        return edited_filepath

    def validate_template(self):
        try:
            doc = Document(self.template_path)
            
            return {
                'valid': True,
                'path': self.template_path,
                'paragraphs_count': len(doc.paragraphs),
                'tables_count': len(doc.tables),
                'fields_found': self.get_template_fields(),
                'error': None
            }
            
        except Exception as e:
            return {
                'valid': False,
                'path': self.template_path,
                'paragraphs_count': 0,
                'tables_count': 0,
                'fields_found': [],
                'error': str(e)
            }
